from docx import Document
import os
from docx import Document
from docx.shared import Inches, Pt
import base64
from telegram_direct import TelegramMessanger
from analytics import calculate_stats

def create_word_stat_document(stats_dict, output_path='output_document.docx'):
    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Statistics Report', level=1)

    # Iterate through the fields in the dictionary
    for field_name, field_stats in stats_dict.items():
        # Add a heading for the field
        doc.add_heading(field_name.capitalize(), level=2)

        # Add a heading for the statistics
        doc.add_heading('Statistics', level=3)

        # Add the statistics as separate messages on separate lines
        for key, value in field_stats.items():
            if key != 'plot_path':
                key_text = key.capitalize()

                # Add the key in bold and the value in regular font
                paragraph = doc.add_paragraph()
                run_key = paragraph.add_run(key_text + ': ')
                run_key.bold = True
                run_value = paragraph.add_run(str(value))
                run_value.bold = False
                run_value.font.size = Pt(14)

        # Add a new line between the statistics and the plot
        doc.add_paragraph()

        # Check if 'plot_path' exists in the field_stats
        if 'plot_path' in field_stats:
            plot_path = field_stats['plot_path']

            # Check if the plot_path is valid
            if os.path.exists(plot_path):
                # Add a heading for the plot
                doc.add_heading('Plot', level=3)

                # Add the plot image to the document
                doc.add_picture(plot_path, width=Inches(6.0))
            else:
                doc.add_paragraph(f"Invalid plot_path: {plot_path}")

    # Save the document
    doc.save(output_path)

messanger = TelegramMessanger("6859309312:AAFo5rGYbvh8cgW4cnH8OW2JNqNckmgqWy8")
stats = calculate_stats("n.andrievskiy@gmail.com", "last_year", "Відчуття дзеркала світу")
doc = create_word_stat_document(stats)
messanger.send_file("output_document.docx", 579467950)