from docx import Document
import os

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import aspose.pdf as ap



def create_word_document(content, output_file_name):


    # Save to Word document
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run(content)
    doc.save(output_file_name)
    print('Word document saved successfully.')


def delete_document(file_path = 'output_report.docx'):
    try:
        os.remove(file_path)
        print(f'File {file_path} deleted successfully.')
    except OSError as e:
        print(f"Error deleting the file: {e}")

